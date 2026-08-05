# -*- coding: utf-8 -*-
from pathlib import Path
import math

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


SCRIPT_DIR = Path(__file__).resolve().parent
ROOT = SCRIPT_DIR.parent if SCRIPT_DIR.name == "06_工具腳本" else SCRIPT_DIR
ASSET_DIR = ROOT / "03_素材與圖片" / "research-assets" / "temple-ai-os"
OUT_DOCX = ROOT / "01_企畫書" / "完整企畫書" / "Temple_AI_OS_競賽企畫書.docx"

LATIN_FONT = "Calibri"
CJK_FONT = "Microsoft JhengHei"
BLACK = RGBColor(29, 33, 39)
MUTED = RGBColor(91, 99, 112)
LINE_GREEN = RGBColor(6, 199, 85)
NAVY = RGBColor(18, 48, 74)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GOLD = RGBColor(154, 108, 27)
LIGHT_FILL = "F4F6F9"
TABLE_HEADER = "E8EEF5"
BORDER = "D8DEE8"


def rgb(hex_color):
    hex_color = hex_color.strip("#")
    return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))


def font_path(bold=False):
    candidates = [
        r"C:\Windows\Fonts\msjhbd.ttc" if bold else r"C:\Windows\Fonts\msjh.ttc",
        r"C:\Windows\Fonts\NotoSansCJK-Regular.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for item in candidates:
        p = Path(item)
        if p.exists():
            return str(p)
    return None


def pil_font(size, bold=False):
    p = font_path(bold)
    if p:
        return ImageFont.truetype(p, size)
    return ImageFont.load_default()


def text_size(draw, text, font):
    box = draw.multiline_textbbox((0, 0), text, font=font, spacing=6, align="center")
    return box[2] - box[0], box[3] - box[1]


def draw_center_text(draw, rect, text, font, fill="#1D2127", spacing=6):
    x1, y1, x2, y2 = rect
    w, h = text_size(draw, text, font)
    draw.multiline_text(
        (x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2),
        text,
        font=font,
        fill=fill,
        spacing=spacing,
        align="center",
    )


def draw_box(draw, rect, title, subtitle=None, fill="#FFFFFF", outline="#C9D2DF", title_fill="#1D2127"):
    draw.rounded_rectangle(rect, radius=24, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = rect
    if subtitle:
        draw_center_text(draw, (x1 + 18, y1 + 18, x2 - 18, y1 + 75), title, pil_font(32, True), title_fill)
        draw_center_text(draw, (x1 + 18, y1 + 80, x2 - 18, y2 - 18), subtitle, pil_font(24), "#3F4856")
    else:
        line_count = title.count("\n") + 1
        title_size = 30 if line_count <= 2 else 24
        draw_center_text(draw, (x1 + 18, y1 + 18, x2 - 18, y2 - 18), title, pil_font(title_size, True), title_fill)


def arrow(draw, start, end, color="#506176", width=5):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    angle = math.atan2(y2 - y1, x2 - x1)
    head_len = 18
    head_angle = math.pi / 7
    p1 = (x2 - head_len * math.cos(angle - head_angle), y2 - head_len * math.sin(angle - head_angle))
    p2 = (x2 - head_len * math.cos(angle + head_angle), y2 - head_len * math.sin(angle + head_angle))
    draw.polygon([end, p1, p2], fill=color)


def create_logo(path):
    img = Image.new("RGBA", (1000, 1000), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    green = "#06C755"
    navy = "#12304A"
    gold = "#A66B1F"
    draw.rounded_rectangle((110, 120, 890, 880), radius=190, fill="#F7FBF8", outline=green, width=18)
    draw.polygon([(260, 430), (500, 255), (740, 430)], fill=navy)
    draw.polygon([(220, 455), (500, 245), (780, 455), (740, 500), (500, 330), (260, 500)], fill=gold)
    draw.rectangle((310, 470, 690, 705), fill="#FFFFFF", outline=navy, width=14)
    for x in (390, 500, 610):
        draw.line((x, 500, x, 700), fill=navy, width=10)
    draw.rounded_rectangle((382, 575, 618, 765), radius=38, fill="#EAF7EE", outline=green, width=12)
    for x in (337, 665):
        draw.line((x, 575, x - 55 if x < 500 else x + 55, 575), fill=green, width=10)
        draw.line((x, 635, x - 55 if x < 500 else x + 55, 635), fill=green, width=10)
        draw.line((x, 695, x - 55 if x < 500 else x + 55, 695), fill=green, width=10)
    draw_center_text(draw, (300, 620, 700, 742), "AI", pil_font(92, True), green)
    draw_center_text(draw, (160, 770, 840, 850), "Temple AI OS", pil_font(50, True), navy)
    img.save(path)


def create_architecture(path):
    img = Image.new("RGB", (1800, 1050), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "Temple AI OS 系統架構", font=pil_font(46, True), fill="#12304A")
    draw.text((72, 105), "LINE 作為入口，AI 與資料平台負責理解、執行與營運分析。", font=pil_font(26), fill="#5B6370")
    boxes = {
        "LINE OA\nRich Menu": (90, 250, 420, 390),
        "Messaging API\nWebhook": (520, 250, 850, 390),
        "FastAPI Backend\n任務與權限": (950, 250, 1280, 390),
        "AI Core\nRAG + Agent": (1380, 250, 1710, 390),
        "LIFF / MINI App\n活動、會員、地圖": (290, 600, 660, 750),
        "PostgreSQL\n使用者與活動資料": (760, 600, 1130, 750),
        "Vector DB\n宮廟知識庫": (1230, 600, 1600, 750),
        "Dashboard\n營運分析與審核": (760, 830, 1130, 970),
    }
    fills = ["#EAF7EE", "#F4F6F9", "#EEF6FF", "#FFF6E6", "#EAF7EE", "#F4F6F9", "#FFF6E6", "#EEF6FF"]
    for (label, rect), fill in zip(boxes.items(), fills):
        draw_box(draw, rect, label, fill=fill, outline="#B8C5D6")
    arrow(draw, (420, 320), (520, 320), "#06A94B")
    arrow(draw, (850, 320), (950, 320), "#506176")
    arrow(draw, (1280, 320), (1380, 320), "#506176")
    arrow(draw, (1115, 390), (970, 600), "#506176")
    arrow(draw, (1490, 390), (1415, 600), "#506176")
    arrow(draw, (1115, 600), (1275, 390), "#506176")
    arrow(draw, (950, 390), (575, 600), "#506176")
    arrow(draw, (945, 750), (945, 830), "#506176")
    img.save(path, quality=95)


def create_rag(path):
    img = Image.new("RGB", (1700, 760), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "AI 核心：RAG + 任務型 Agent", font=pil_font(42, True), fill="#12304A")
    steps = [
        ("宮廟資料", "官網、公告、訪談、文化部資料"),
        ("整理審核", "人工校對、分類、版本管理"),
        ("向量檢索", "依問題找最相關段落"),
        ("LLM 生成", "只根據可信內容回答"),
        ("任務執行", "查活動、報名、通知、摘要"),
    ]
    x = 70
    for i, (title, sub) in enumerate(steps):
        rect = (x, 220, x + 270, 390)
        fill = ["#EAF7EE", "#F4F6F9", "#FFF6E6", "#EEF6FF", "#EAF7EE"][i]
        draw_box(draw, rect, title, sub, fill=fill, outline="#B8C5D6")
        if i < len(steps) - 1:
            arrow(draw, (x + 270, 305), (x + 335, 305), "#506176")
        x += 335
    draw.rounded_rectangle((185, 520, 1515, 645), radius=22, fill="#FAFBFC", outline="#D8DEE8", width=3)
    draw.text((225, 545), "內容治理原則", font=pil_font(30, True), fill="#12304A")
    draw.text(
        (225, 590),
        "不扮演神明、不預測命運、不自創宗教解釋；涉及文化與儀式內容必須可追溯來源並保留人工審核。",
        font=pil_font(25),
        fill="#3F4856",
    )
    img.save(path, quality=95)


def create_journey(path):
    img = Image.new("RGB", (1700, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "核心使用者旅程", font=pil_font(42, True), fill="#12304A")
    lanes = [
        ("新信眾 / 遊客", ["掃 QR/NFC", "加入 LINE OA", "AI 導覽", "LIFF 報名", "收到提醒"]),
        ("固定信眾", ["Rich Menu", "查活動", "會員中心", "參與任務", "累積互動"]),
        ("宮廟管理者", ["登入後台", "新增活動", "審核知識", "查看分析", "推播通知"]),
    ]
    y = 170
    for lane, steps in lanes:
        draw.text((80, y + 46), lane, font=pil_font(30, True), fill="#12304A")
        x = 330
        for i, step in enumerate(steps):
            rect = (x, y, x + 210, y + 110)
            draw_box(draw, rect, step, fill=["#EAF7EE", "#F4F6F9", "#FFF6E6", "#EEF6FF", "#EAF7EE"][i], outline="#B8C5D6")
            if i < len(steps) - 1:
                arrow(draw, (x + 210, y + 55), (x + 270, y + 55), "#506176", 4)
            x += 270
        y += 230
    img.save(path, quality=95)


def create_er(path):
    img = Image.new("RGB", (1700, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "資料模型概念 ER Diagram", font=pil_font(42, True), fill="#12304A")
    entities = {
        "User\nline_user_id\nprofile\nconsent": (90, 190, 410, 350),
        "Message\nuser_id\nintent\ncontent": (510, 190, 830, 350),
        "Event\ntitle\ntime\ncapacity": (930, 190, 1250, 350),
        "Registration\nuser_id\nevent_id\nstatus": (510, 520, 830, 680),
        "Knowledge\nsource\nchunk\nreviewed": (930, 520, 1250, 680),
        "Point\nuser_id\nreason\namount": (90, 520, 410, 680),
        "Dashboard\nmetrics\nsegments\nreports": (1350, 360, 1630, 520),
    }
    for label, rect in entities.items():
        draw_box(draw, rect, label, fill="#F4F6F9", outline="#B8C5D6")
    arrow(draw, (410, 270), (510, 270), "#506176")
    arrow(draw, (670, 350), (670, 520), "#506176")
    arrow(draw, (830, 600), (930, 300), "#506176")
    arrow(draw, (410, 600), (510, 600), "#506176")
    arrow(draw, (1250, 270), (1410, 360), "#506176")
    arrow(draw, (1250, 600), (1410, 520), "#506176")
    arrow(draw, (830, 600), (1350, 440), "#506176")
    img.save(path, quality=95)


def set_run_font(run, name=LATIN_FONT, east_asia=CJK_FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size=None, color=None, bold=None, italic=None):
    style.font.name = LATIN_FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def paragraph_spacing(paragraph, before=0, after=8, line=1.333):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_para(doc, text="", style=None, bold=False, color=BLACK, size=None, align=None, after=8, before=0, italic=False):
    p = doc.add_paragraph(style=style)
    paragraph_spacing(p, before=before, after=after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_kv_para(doc, label, value):
    p = doc.add_paragraph()
    paragraph_spacing(p, after=4, line=1.2)
    r = p.add_run(f"{label}：")
    set_run_font(r, size=10.5, bold=True, color=NAVY)
    r = p.add_run(value)
    set_run_font(r, size=10.5, color=BLACK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        paragraph_spacing(p, before=18, after=10, line=1.15)
    elif level == 2:
        paragraph_spacing(p, before=12, after=6, line=1.15)
    else:
        paragraph_spacing(p, before=8, after=4, line=1.15)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    paragraph_spacing(p, after=4, line=1.208)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=BLACK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    paragraph_spacing(p, after=4, line=1.208)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=BLACK)
    return p


def add_callout(doc, title, body, fill=LIGHT_FILL, trailing=True):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], indent=120, borders=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    paragraph_spacing(p, after=4, line=1.15)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    paragraph_spacing(p2, after=0, line=1.25)
    r = p2.add_run(body)
    set_run_font(r, size=10.5, color=BLACK)
    if trailing:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:tcMar"))
    if old is not None:
        tc_pr.remove(old)
    tc_mar = OxmlElement("w:tcMar")
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_cell_text(cell, text, bold=False, color=BLACK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    paragraph_spacing(p, after=0, line=1.15)
    p.alignment = align
    for i, part in enumerate(str(text).split("\n")):
        if i:
            p.add_run().add_break()
        r = p.add_run(part)
        set_run_font(r, size=size, color=color, bold=bold)


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths, indent=120, borders=True):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    for grid in tbl.xpath("./w:tblGrid"):
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(idx, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
    if borders:
        set_table_borders(table)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def add_table(doc, headers, rows, widths, header_fill=TABLE_HEADER):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_text(cell, h, bold=True, color=NAVY, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    keep_row_together(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        keep_row_together(row)
        cells = row.cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def add_picture(doc, path, caption):
    p = doc.add_paragraph()
    paragraph_spacing(p, before=8, after=4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.3))
    c = doc.add_paragraph()
    paragraph_spacing(c, before=0, after=8, line=1.15)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    set_run_font(r, size=9, color=MUTED, italic=True)


def add_page_number(paragraph):
    def field(instr):
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instr
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr_text)
        run._r.append(sep)
        run._r.append(end)
        return run

    paragraph.add_run("第 ")
    field("PAGE")
    paragraph.add_run(" / ")
    field("NUMPAGES")
    paragraph.add_run(" 頁")
    for r in paragraph.runs:
        set_run_font(r, size=9, color=MUTED)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], 11, BLACK)
    styles["Normal"].paragraph_format.space_after = Pt(8)
    styles["Normal"].paragraph_format.line_spacing = 1.333
    set_style_font(styles["Heading 1"], 16, BLUE, True)
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    set_style_font(styles["Heading 2"], 13, BLUE, True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    set_style_font(styles["Heading 3"], 12, DARK_BLUE, True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    set_style_font(styles["List Bullet"], 10.5, BLACK)
    set_style_font(styles["List Number"], 10.5, BLACK)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("Temple AI OS | 2026 LINE AI 創新創業競賽")
    set_run_font(r, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)


def cover(doc, logo):
    add_para(doc, "2026 LINE AI 創新創業競賽提案書", bold=True, color=MUTED, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.add_run().add_picture(str(logo), width=Inches(1.35))
    add_para(doc, "Temple AI OS", bold=True, color=NAVY, size=30, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "AI 驅動的宮廟智慧社群營運平台", color=BLACK, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(
        doc,
        "以 LINE OA、Messaging API、LIFF / MINI App 與 AI RAG Agent，打造宮廟文化導覽、活動服務、會員經營與營運分析的一站式平台。",
        color=MUTED,
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=20,
    )
    add_table(
        doc,
        ["項目", "內容"],
        [
            ["團隊名稱", "[請填入團隊名稱]"],
            ["參賽題目", "智慧宮廟 AI 服務入口與數位營運平台"],
            ["版本", "企畫書 + 產品規格書 + 開發企劃書 + Demo 製作指南"],
            ["整理日期", "2026-08-02"],
            ["聯絡資訊", "[請填入聯絡人、Email、電話]"],
        ],
        [1900, 7460],
    )
    add_callout(
        doc,
        "一句話提案",
        "Temple AI OS 不是 AI 算命或單純聊天機器人，而是把 LINE 變成宮廟服務入口，讓信眾可問、可報名、可收到通知，讓管理者可更新內容、管理活動並看見營運數據。",
        fill="EAF7EE",
    )
    doc.add_page_break()


def toc(doc):
    add_heading(doc, "目錄與文件定位", 1)
    add_para(
        doc,
        "本文件將原始目錄整理成可交付的競賽企畫書。前半段回答評審最在意的市場、創新、LINE 生態整合與商業模式；後半段整理開發規格、Demo 腳本與資料蒐集清單，方便後續拆成 10 頁簡報與開發任務。",
    )
    items = [
        "一、競賽對齊與專案摘要",
        "二、市場痛點、使用者與競品分析",
        "三、解決方案與產品功能規格",
        "四、LINE 與 AI 技術設計",
        "五、資料庫、API、Dashboard 與 UI/UX",
        "六、Demo 製作指南、開發流程與測試計畫",
        "七、風險管理、商業模式與資料來源",
    ]
    for item in items:
        add_bullet(doc, item)
    doc.add_page_break()


def section_competition(doc):
    add_heading(doc, "一、競賽對齊與專案摘要", 1)
    add_heading(doc, "1.1 Executive Summary", 2)
    add_para(
        doc,
        "Temple AI OS 以智慧宮廟為垂直場景，將 LINE OA 作為信眾入口，透過 Messaging API 接收對話、LIFF / MINI App 承載會員與活動流程，並以 RAG 與任務型 Agent 提供可追溯、可審核的文化導覽與服務協助。",
    )
    add_para(
        doc,
        "此提案鎖定三個問題：傳統宮廟資訊分散、年輕族群參與門檻高、管理者缺乏低成本的數位營運工具。平台用 LINE 降低使用門檻，用 AI 降低服務與內容維護成本，用 Dashboard 讓宮廟看見活動、問題與會員互動成效。",
    )
    add_callout(
        doc,
        "核心差異",
        "評審不會只看到一個會聊天的 Bot，而會看到完整 LINE 生態服務：OA 導流、Rich Menu 導航、Chatbot 問答、LIFF / MINI App 報名與會員頁、Flex Message 結果卡、後台分析與未來 LINE Touch / Pay 擴充。",
    )
    add_heading(doc, "1.2 競賽規則與評分對應", 2)
    add_table(
        doc,
        ["競賽重點", "企畫書對應設計"],
        [
            ["AI 服務互動展示", "AI 導覽、文化抽籤、活動查詢、報名提醒、常見問題摘要。"],
            ["目標族群 / 市場需求", "以新信眾、固定信眾、宮廟管理者與周邊商圈作為四類使用者。"],
            ["LINE App 應用創意性", "整合 OA、Rich Menu、Messaging API、Flex Message、LIFF / MINI App。"],
            ["營運模式 / 商業模式", "SaaS 月費、導入費、祭典活動服務與周邊商圈加值模組。"],
            ["LINE 生態體系結合度", "以 LINE 作為入口、通知、會員識別、互動與服務閉環。"],
        ],
        [2100, 7260],
    )
    add_table(
        doc,
        ["交件 / 賽程資訊", "企畫書使用方式"],
        [
            ["投稿簡報 PDF 10 頁內，不含封面與封底", "本文件可濃縮為 10 頁簡報：痛點、方案、Demo、商模與技術。"],
            ["Demo 影片 YouTube Link 3 分鐘內", "第六章提供 3 分鐘投稿影片腳本與鏡頭順序。"],
            ["決選簡報 5 分鐘、問答 5 分鐘", "第六章另提供 5 分鐘現場 Demo 版本。"],
            ["投稿截止 2026-09-24 17:00 Asia/Taipei", "開發時程以 8 月原型、9 月簡報與影片收斂為主。"],
        ],
        [2600, 6760],
    )
    add_heading(doc, "1.3 專案品牌與視覺方向", 2)
    add_para(doc, "建議正式名稱使用 Temple AI OS，避免以「AI 求籤機器人」作為主標，因為後者容易被評審理解為單點娛樂功能，無法呈現營運平台的完整性。")
    add_bullet(doc, "主標：Temple AI OS")
    add_bullet(doc, "副標：AI 驅動的宮廟智慧社群營運平台")
    add_bullet(doc, "Logo 元素：廟宇屋簷、香火意象、AI 晶片線路、LINE 綠色入口感。")
    add_bullet(doc, "Slogan：讓宮廟服務走進 LINE，讓文化知識被理解、被參與、被延續。")


def section_market(doc):
    add_heading(doc, "二、市場痛點、使用者與競品分析", 1)
    add_heading(doc, "2.1 市場痛點", 2)
    add_para(
        doc,
        "台灣宮廟具備高度在地性與社群影響力，但服務資訊常散落在公告欄、Facebook、官網、口耳相傳或現場人員說明中。當年輕使用者不熟悉參拜流程、祭典文化或報名方式時，通常需要現場詢問；當宮廟管理者要處理活動、通知與問題回覆時，也缺乏整合的數位營運工具。",
    )
    add_table(
        doc,
        ["痛點", "現況", "Temple AI OS 回應"],
        [
            ["資訊分散", "官網、社群與紙本公告不一致，信眾難以找到最新資訊。", "由後台統一維護活動與 FAQ，AI 回答必須引用知識庫。"],
            ["參與門檻", "新信眾不熟悉儀式、祭典與動線，年輕族群更偏好即時互動。", "用 LINE 對話與 LIFF 導覽降低首次使用成本。"],
            ["行政負擔", "活動報名、問題回覆、通知與統計多靠人工。", "Agent 協助查詢、報名、提醒、摘要常見問題。"],
            ["營運不可見", "管理者難以知道熱門問題、活動轉換率與會員活躍度。", "Dashboard 顯示活動、AI 問答、會員與流量分析。"],
        ],
        [1700, 3600, 4060],
    )
    add_heading(doc, "2.2 目標使用者", 2)
    add_table(
        doc,
        ["使用者", "主要需求", "MVP 場景"],
        [
            ["新信眾 / 遊客", "快速理解宮廟文化、參拜流程與活動資訊。", "掃 QR 加入 LINE，問 AI「第一次來要怎麼拜？」"],
            ["固定信眾", "接收活動提醒、報名服務、查看會員與互動紀錄。", "Rich Menu 進入會員中心並完成祭典報名。"],
            ["宮廟管理者", "降低客服與行政負擔，維護資料並查看成效。", "後台新增活動、審核知識、查看熱門問題。"],
            ["周邊攤商 / 合作單位", "祭典期間獲得導流、優惠曝光與人流資訊。", "MINI App 地圖呈現攤商資訊與活動動線。"],
        ],
        [1900, 3800, 3660],
    )
    add_heading(doc, "2.3 競品與替代方案", 2)
    add_para(doc, "競品分析不主打「別人不好」，而是指出現有服務多半是單點工具，缺少資訊、互動、會員、活動與營運分析的整合。")
    add_table(
        doc,
        ["類型", "常見能力", "限制", "Temple AI OS 差異"],
        [
            ["宮廟官網", "介紹沿革、活動公告、交通資訊。", "更新成本高，互動弱。", "以 LINE 作日常入口，後台更新即可供 AI 與 LIFF 使用。"],
            ["Facebook 粉專", "公告與照片分享。", "資訊排序與搜尋弱，會員流程不足。", "提供可查詢、可報名、可通知的服務閉環。"],
            ["一般 LINE OA", "群發、關鍵字、自動回覆。", "通常缺少 AI、資料庫與後台分析。", "把 OA 接到 Webhook、AI、會員與活動系統。"],
            ["線上求籤網站", "互動娛樂。", "容易落入算命與命運預測風險。", "定位為文化互動體驗，提供籤詩背景與正向解說。"],
        ],
        [1700, 2500, 2500, 2660],
    )
    add_heading(doc, "2.4 創新性", 2)
    add_bullet(doc, "不是單純聊天：AI 只是一個入口，真正價值在活動、會員、通知、分析與服務流程。")
    add_bullet(doc, "不是一般 GPT：回覆由宮廟知識庫、活動資料與審核流程約束。")
    add_bullet(doc, "不是一次性 Demo：Dashboard 與資料模型支撐後續營運。")
    add_bullet(doc, "不是宗教預測：所有 AI 文化互動都避免斷言吉凶與個人命運。")


def section_solution(doc, architecture, journey):
    add_heading(doc, "三、解決方案與產品功能規格", 1)
    add_picture(doc, architecture, "圖 1：Temple AI OS 系統架構圖")
    add_heading(doc, "3.1 三大系統", 2)
    add_table(
        doc,
        ["系統", "包含功能", "Demo 必須呈現"],
        [
            ["LINE 智慧服務", "OA、Rich Menu、Messaging API、Flex Message、LIFF / MINI App。", "使用者從 LINE 進入，完成問答與報名。"],
            ["AI 核心", "RAG 知識庫、任務型 Agent、摘要與推薦。", "AI 回答宮廟文化與活動問題，並協助下一步操作。"],
            ["管理平台", "活動管理、知識審核、會員 CRM、分析 Dashboard。", "管理者可看到報名數、熱門問題與推播入口。"],
        ],
        [1900, 4300, 3160],
    )
    add_heading(doc, "3.2 MVP 功能範圍", 2)
    add_table(
        doc,
        ["功能", "目的", "輸入", "輸出 / 成果"],
        [
            ["AI 宮廟導覽", "回答沿革、神明文化、參拜流程與活動資訊。", "LINE 文字問題。", "文字回覆與 Flex Message 引導卡。"],
            ["AI 文化抽籤", "提高互動，但定位為文化體驗。", "主題選擇或心情選項。", "籤詩、文化背景、正向建議與注意事項。"],
            ["活動中心", "處理祭典、導覽、志工與報名流程。", "LIFF 表單與會員資料。", "報名紀錄、確認訊息、提醒通知。"],
            ["會員中心", "建立可識別、可互動的信眾服務。", "LINE Login / LIFF Profile。", "會員頁、互動紀錄、服務入口。"],
            ["後台 Dashboard", "讓管理者看見營運效果。", "活動、問答、會員與點擊資料。", "報名率、熱門問題、活躍人數與內容缺口。"],
        ],
        [1600, 3000, 2300, 2460],
    )
    add_heading(doc, "3.3 加分功能與商轉擴充", 2)
    add_bullet(doc, "LINE Touch / NFC：祭典現場碰觸入口，開啟導覽、報到或集章。")
    add_bullet(doc, "LINE Pay：報名費、服務費、周邊商品或香油錢付款；MVP 可先模擬付款流程。")
    add_bullet(doc, "社群摘要：整理 LINE 群組或社群回饋，協助管理者掌握問題。")
    add_bullet(doc, "攤商服務：祭典地圖、優惠券、導流與點擊分析。")
    add_bullet(doc, "Beacon / 人流資料：未來可結合現場感測或人工回報，做壅塞提醒。")
    add_picture(doc, journey, "圖 2：新信眾、固定信眾與管理者的核心旅程")


def section_tech(doc, rag):
    add_heading(doc, "四、LINE 與 AI 技術設計", 1)
    add_heading(doc, "4.1 LINE 技術配置", 2)
    add_table(
        doc,
        ["LINE 技術", "用途", "Demo 實作重點"],
        [
            ["Official Account / OA", "官方帳號入口、好友經營、Rich Menu 與推播。", "建立服務入口，讓評審從 LINE 體驗。"],
            ["Messaging API", "接收 Webhook 事件並回覆訊息。", "/webhook 收訊息，呼叫 AI 後以 Reply API 回覆。"],
            ["Flex Message", "以卡片呈現活動、籤詩、導覽結果。", "讓 Demo 視覺比純文字更像正式服務。"],
            ["LIFF / MINI App", "在 LINE 內開啟會員、表單、地圖與活動頁。", "活動報名、會員中心、祭典地圖。"],
            ["LINE Login", "識別會員並串接個人化資料。", "取得 LINE user ID 並建立會員紀錄。"],
            ["LINE Pay / Touch", "付款與現場 OMO 入口。", "列為加分模組，Demo 可用 QR/NFC 模擬。"],
        ],
        [1900, 3500, 3960],
    )
    add_heading(doc, "4.2 AI 技術設計", 2)
    add_picture(doc, rag, "圖 3：RAG 與任務型 Agent 流程")
    add_para(
        doc,
        "AI 核心由兩層組成。第一層是 RAG 知識問答，用於宮廟沿革、神明文化、參拜流程與公告內容，避免模型自由發揮。第二層是任務型 Agent，用於查活動、建立報名、提醒、摘要常見問題與推薦下一步服務。",
    )
    add_table(
        doc,
        ["AI 模組", "資料來源", "控制方式"],
        [
            ["宮廟文化問答", "官網、文化部資料、國家文化記憶庫、現場訪談。", "回答需依檢索內容生成，無來源時轉人工或回覆無法確認。"],
            ["活動問答與報名", "管理後台活動表、名額、時間與地點。", "Agent 只可呼叫白名單 API，例如查詢、建立報名、取消報名。"],
            ["AI 文化抽籤", "審核過的籤詩文本與文化解說。", "避免預測命運，用文化背景與正向提醒取代吉凶斷言。"],
            ["營運摘要", "匿名化訊息、熱門問題、活動紀錄。", "移除敏感個資，只輸出趨勢、常見問題與內容缺口。"],
        ],
        [1800, 3100, 4460],
    )
    add_heading(doc, "4.3 推薦技術選型", 2)
    add_table(
        doc,
        ["層級", "建議技術", "原因"],
        [
            ["Frontend", "Next.js + LIFF SDK", "可同時做官網、LIFF 頁、後台與未來 MINI App 介面。"],
            ["Backend", "FastAPI", "Python 生態適合 AI、Webhook、資料處理與快速 Demo。"],
            ["Database", "PostgreSQL；MVP 可 SQLite", "支援會員、活動、報名、訊息與分析資料。"],
            ["Vector DB", "Chroma；商轉可 Pinecone / pgvector", "MVP 輕量，商轉可擴充。"],
            ["AI Provider", "OpenAI API / Gemini", "依成本、延遲、繁中表現與團隊熟悉度選擇。"],
            ["部署", "Render / Railway / Vercel / Cloud Run", "需可提供 HTTPS Webhook 與穩定 Demo URL。"],
        ],
        [1700, 3300, 4360],
    )


def section_specs(doc, er):
    add_heading(doc, "五、資料庫、API、Dashboard 與 UI/UX", 1)
    add_heading(doc, "5.1 API 設計", 2)
    add_table(
        doc,
        ["API", "方法", "用途", "Demo 狀態"],
        [
            ["/webhook", "POST", "接收 LINE Messaging API 事件。", "必做"],
            ["/chat", "POST", "呼叫 RAG / Agent 生成回覆。", "必做"],
            ["/users/me", "GET", "取得 LIFF / LINE Login 綁定會員。", "必做"],
            ["/events", "GET / POST", "活動列表與後台新增活動。", "必做"],
            ["/registrations", "POST", "建立活動報名紀錄。", "必做"],
            ["/knowledge", "GET / POST", "上傳與審核宮廟知識資料。", "加分"],
            ["/analytics/summary", "GET", "Dashboard 指標與熱門問題。", "必做"],
        ],
        [1700, 1300, 4300, 2060],
    )
    add_heading(doc, "5.2 資料庫設計", 2)
    add_picture(doc, er, "圖 4：MVP 資料模型概念")
    add_table(
        doc,
        ["資料表", "核心欄位", "用途"],
        [
            ["User", "id, line_user_id, display_name, consent_at", "會員識別、個人化與通知權限。"],
            ["Message", "id, user_id, role, content, intent, created_at", "AI 問答紀錄、熱門問題分析。"],
            ["Event", "id, title, start_at, location, capacity", "祭典、導覽、志工與活動資料。"],
            ["Registration", "id, user_id, event_id, status", "活動報名與提醒通知。"],
            ["Knowledge", "id, source, category, chunk, reviewed_by", "RAG 檢索資料與內容治理。"],
            ["Point", "id, user_id, reason, amount", "互動任務、集章或會員活躍度。"],
        ],
        [1700, 3600, 4060],
    )
    add_heading(doc, "5.3 Dashboard 指標", 2)
    add_bullet(doc, "使用者分析：好友數、活躍人數、新增會員、回訪率。")
    add_bullet(doc, "活動分析：瀏覽、報名、取消、到場與提醒點擊。")
    add_bullet(doc, "AI 分析：熱門問題、無法回答問題、知識庫缺口、常見意圖。")
    add_bullet(doc, "內容審核：新增資料、來源、審核狀態與最後更新時間。")
    add_heading(doc, "5.4 UI/UX 設計原則", 2)
    add_table(
        doc,
        ["介面", "設計重點", "要避免"],
        [
            ["LINE Rich Menu", "3-4 個主入口：AI 導覽、活動中心、會員中心、管理入口。", "入口太多、文字太小、功能命名不清。"],
            ["Flex Message", "活動卡、導覽卡、報名成功卡、文化抽籤卡。", "只有長文字，缺少下一步按鈕。"],
            ["LIFF / MINI App", "手機優先，保留 LINE 內安全區，表單短而清楚。", "像一般網頁，不像 LINE 內服務。"],
            ["Dashboard", "給管理者掃描重點，首頁呈現指標與待處理項目。", "堆滿圖表但無法行動。"],
        ],
        [1800, 3900, 3660],
    )


def section_demo_dev(doc):
    add_heading(doc, "六、Demo 製作指南、開發流程與測試計畫", 1)
    add_heading(doc, "6.1 三分鐘投稿影片腳本", 2)
    add_table(
        doc,
        ["時間", "畫面", "旁白重點"],
        [
            ["0:00-0:25", "宮廟公告、排隊、FB 與紙本資訊切換。", "傳統宮廟資訊分散，年輕使用者不知道怎麼參與。"],
            ["0:25-0:55", "掃 QR 加入 LINE OA，打開 Rich Menu。", "Temple AI OS 把服務入口放進使用者每天使用的 LINE。"],
            ["0:55-1:35", "問 AI 參拜流程與活動時間，顯示 Flex Message。", "AI 不是自由算命，而是根據審核知識庫回答。"],
            ["1:35-2:15", "開啟 LIFF 活動中心，完成報名。", "聊天之外，使用者可在 LINE 內完成服務流程。"],
            ["2:15-2:45", "後台 Dashboard 顯示報名、熱門問題與知識缺口。", "宮廟管理者能看見營運數據，降低行政負擔。"],
            ["2:45-3:00", "商業模式與未來擴充畫面。", "SaaS 導入宮廟，擴充 LINE Touch、Pay、攤商與祭典服務。"],
        ],
        [1500, 4000, 3860],
    )
    add_heading(doc, "6.2 五分鐘決選 Demo 腳本", 2)
    for step in [
        "0:00-1:00 痛點與目標族群：說明宮廟資訊、參與與管理問題。",
        "1:00-2:00 LINE 入口：展示 OA、Rich Menu、QR / NFC 入口概念。",
        "2:00-3:00 AI 互動：問答、文化抽籤、RAG 來源與安全邊界。",
        "3:00-4:00 LIFF / MINI App 流程：會員、活動報名、提醒通知。",
        "4:00-5:00 Dashboard 與商業模式：數據、導入對象、月費與擴充。",
    ]:
        add_number(doc, step)
    add_heading(doc, "6.3 開發流程", 2)
    add_table(
        doc,
        ["Phase", "目標", "產出"],
        [
            ["Phase 1", "LINE OA + Webhook 原型", "能接收訊息、回覆 Flex Message、展示 Rich Menu。"],
            ["Phase 2", "AI RAG 問答", "宮廟知識庫、檢索、回答限制與 fallback。"],
            ["Phase 3", "LIFF 活動與會員", "會員頁、活動列表、報名表、成功通知。"],
            ["Phase 4", "Dashboard", "活動、問答、會員與內容審核指標。"],
            ["Phase 5", "Demo 收斂", "影片腳本、假資料、測試劇本、簡報圖表。"],
        ],
        [1600, 3000, 4760],
    )
    add_heading(doc, "6.4 測試計畫", 2)
    add_table(
        doc,
        ["測試面向", "驗收標準"],
        [
            ["AI 回覆正確率", "宮廟文化與活動問題至少準備 30 題測試集；無資料時不硬答。"],
            ["Webhook 穩定性", "LINE 訊息可在 3 秒內回應基本接收，較慢 AI 任務有 loading 或延遲策略。"],
            ["LIFF 流程", "手機畫面可完成登入、報名與成功頁，不需離開 LINE。"],
            ["資料一致性", "後台新增活動後，AI 與活動頁都讀到同一份資料。"],
            ["Demo 腳本", "3 分鐘影片可完整展示痛點、服務、AI、報名與 Dashboard。"],
        ],
        [2400, 6960],
    )


def section_risk_business_sources(doc):
    add_heading(doc, "七、風險管理、商業模式與資料來源", 1)
    add_heading(doc, "7.1 風險管理", 2)
    add_table(
        doc,
        ["風險", "影響", "處理方式"],
        [
            ["宗教內容錯誤", "傷害信任與文化尊重。", "建立來源、審核者與版本紀錄；敏感內容轉人工審核。"],
            ["AI 幻覺", "回答不存在的活動或錯誤文化解釋。", "RAG 限制、白名單工具、無來源時拒答或請使用者洽廟方。"],
            ["個資與宗教敏感資料", "違反使用者信任與法規要求。", "只收必要資料，清楚告知用途，分析資料匿名化。"],
            ["功能範圍過大", "Demo 做不完、主軸分散。", "MVP 聚焦 LINE 入口、AI 問答、活動報名、Dashboard。"],
            ["LINE Touch / Pay 實作門檻", "正式串接需申請或商業條件。", "比賽用 QR/NFC 與假付款流程展示概念，列為商轉加分。"],
        ],
        [2200, 2900, 4260],
    )
    add_heading(doc, "7.2 商業模式", 2)
    add_table(
        doc,
        ["商業模式構面", "規劃"],
        [
            ["客戶", "中大型宮廟、地方文化組織、祭典主辦單位、周邊商圈。"],
            ["價值主張", "降低服務與行政成本，提高年輕族群參與，沉澱可用的文化與營運資料。"],
            ["收入", "SaaS 月費、導入建置費、祭典短期專案、攤商曝光與資料分析加值。"],
            ["成本", "雲端主機、LLM API、向量資料庫、LINE 訊息費用、內容整理與客服。"],
            ["通路", "LINE 技術社群、地方創生單位、文化局合作、宮廟聯誼組織。"],
        ],
        [2300, 7060],
    )
    add_heading(doc, "7.3 10 頁簡報建議結構", 2)
    add_table(
        doc,
        ["頁次", "主題", "內容"],
        [
            ["1", "題目與一句話提案", "Temple AI OS、目標場景、核心價值。"],
            ["2", "痛點與使用者", "三大痛點與四類使用者。"],
            ["3", "解決方案", "LINE + AI + Dashboard 三大系統。"],
            ["4", "Demo 使用者旅程", "掃 QR、問 AI、報名、收到提醒。"],
            ["5", "LINE 生態整合", "OA、Messaging API、LIFF / MINI App、Flex。"],
            ["6", "AI 技術", "RAG、Agent、內容治理與安全邊界。"],
            ["7", "市場與競品", "現有服務缺口與 Temple AI OS 差異。"],
            ["8", "商業模式", "客戶、收入、成本、導入方式。"],
            ["9", "開發時程與測試", "MVP、加分功能、測試計畫。"],
            ["10", "結語與未來擴充", "LINE Touch、Pay、攤商、祭典人流。"],
        ],
        [900, 2600, 5860],
    )
    add_heading(doc, "7.4 資料蒐集與參考來源", 2)
    add_para(doc, "後續正式報名文件建議持續查核下列來源，並在簡報末頁保留簡短來源清單。")
    for source in [
        "競賽頁：https://contest.bhuntr.com/tw/rst0pkz3f3yoxx89ry/home/",
        "競賽說明：https://contest.bhuntr.com/tw/rst0pkz3f3yoxx89ry/details/",
        "LINE Developers：https://developers.line.biz/",
        "Messaging API：https://developers.line.biz/en/docs/messaging-api/",
        "LIFF：https://developers.line.biz/en/docs/liff/",
        "LINE MINI App：https://developers.line.biz/en/docs/line-mini-app/",
        "文化部資料庫、國家文化記憶庫、內政部宗教資訊、各宮廟官方資料。",
        "Google Scholar 關鍵字：台灣民間信仰、宮廟數位轉型、宗教文化保存。",
    ]:
        add_bullet(doc, source)
    add_heading(doc, "7.5 最終範圍建議", 2)
    add_para(
        doc,
        "為了在競賽期限內完成高品質 Demo，建議必做 LINE OA、AI 問答、宮廟知識庫、AI 文化抽籤、活動管理、LIFF 會員頁、官網雛形與管理後台；LINE Pay、NFC、攤商服務與人流分析作為加分或簡報未來規劃。",
        after=0,
    )


def build():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    logo = ASSET_DIR / "temple_ai_os_logo.png"
    architecture = ASSET_DIR / "architecture.png"
    rag = ASSET_DIR / "rag_agent.png"
    journey = ASSET_DIR / "journey.png"
    er = ASSET_DIR / "er_diagram.png"
    create_logo(logo)
    create_architecture(architecture)
    create_rag(rag)
    create_journey(journey)
    create_er(er)

    doc = Document()
    configure_doc(doc)
    props = doc.core_properties
    props.title = "Temple AI OS 競賽企畫書"
    props.subject = "2026 LINE AI 創新創業競賽 - 智慧宮廟提案"
    props.author = "Codex"
    props.keywords = "LINE, AI, Temple AI OS, 智慧宮廟, LIFF, Messaging API, RAG"

    cover(doc, logo)
    toc(doc)
    section_competition(doc)
    section_market(doc)
    section_solution(doc, architecture, journey)
    section_tech(doc, rag)
    section_specs(doc, er)
    section_demo_dev(doc)
    section_risk_business_sources(doc)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
